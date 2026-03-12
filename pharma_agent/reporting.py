from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from pharma_agent.models import ResearchResult


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relation_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_label_and_link(document: Document, label: str, url: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}: ")
    if url and url != "Not found":
        add_hyperlink(paragraph, url, url)
    else:
        paragraph.add_run("Not found")


def build_report(result: ResearchResult, output_file: str | Path) -> Path:
    document = Document()
    document.add_heading(result.title, level=1)
    document.add_paragraph(f"Location: {result.query}")
    document.add_paragraph(f"Companies found: {len(result.companies)}")

    for index, record in enumerate(result.companies, start=1):
        document.add_heading(f"{index}. {record.name}", level=2)
        document.add_paragraph(f"Location: {record.location}")
        add_label_and_link(document, "Website", record.website)
        document.add_paragraph(f"Email: {record.email}")
        document.add_paragraph(f"Phone: {record.phone}")
        document.add_paragraph(f"Products / Specialization: {record.products}")
        document.add_paragraph(f"Description: {record.description}")

        if record.jobs:
            document.add_paragraph("Open Roles:")
            for job in record.jobs[:5]:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.add_run(f"{job.title} | {job.location} | Apply: ")
                if job.apply_link and job.apply_link != "Not found":
                    add_hyperlink(paragraph, job.apply_link, job.apply_link)
                else:
                    paragraph.add_run("Not found")

    path = Path(output_file)
    document.save(path)
    return path
